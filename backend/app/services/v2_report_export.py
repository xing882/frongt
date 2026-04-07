from __future__ import annotations

import io
import json
from datetime import datetime
from typing import Any

from app.services import stats_service
from app.services.v2_service import ops_indicators, ops_suggestions

try:
    from docx import Document

    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    _HAS_REPORTLAB = True
except ImportError:
    _HAS_REPORTLAB = False


def _report_context(kind: str, building_id: str | None) -> dict[str, Any]:
    period = stats_service.period_summary(building_id, None, None)
    ind = ops_indicators(building_id)
    sug = ops_suggestions(building_id)
    return {
        "kind": kind,
        "building_id": building_id,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "period_summary": period,
        "indicators": ind.get("indicators", {}),
        "suggestions": sug.get("items", []),
    }


def build_report_bytes(kind: str, file_format: str, building_id: str | None = None) -> tuple[bytes, str, str]:
    """
    Returns (body, filename, media_type).
    """
    ctx = _report_context(kind, building_id)
    title = "建筑能源运营优化报告" if kind == "operations" else "建筑能源 ESG 专项报告"

    if file_format == "word":
        if not _HAS_DOCX:
            txt = json.dumps(ctx, ensure_ascii=False, indent=2).encode("utf-8")
            return txt, f"{kind}_report_fallback.txt", "text/plain; charset=utf-8"
        bio = io.BytesIO()
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"生成时间：{ctx['generated_at']}")
        doc.add_paragraph(f"建筑范围：{building_id or '全部（未筛选）'}")

        doc.add_heading("一、时段能耗汇总（V0.5 基础数据）", level=1)
        doc.add_paragraph(json.dumps(ctx["period_summary"], ensure_ascii=False, indent=2))

        doc.add_heading("二、运营核心指标（V2.0）", level=1)
        doc.add_paragraph(json.dumps(ctx["indicators"], ensure_ascii=False, indent=2))

        if kind == "esg":
            doc.add_heading("三、ESG 维度说明（演示）", level=1)
            doc.add_paragraph(
                "环境：碳排与能耗强度基于现有小时电耗数据推演；"
                "社会：舒适度以温度/湿度均值表征；"
                "治理：设备与工单健康度由运营指标近似。"
            )

        doc.add_heading("三、优化建议清单" if kind == "operations" else "四、优化建议清单", level=1)
        for i, it in enumerate(ctx["suggestions"], 1):
            doc.add_paragraph(f"{i}. [{it.get('priority', '-')}] {it.get('title', '')}")

        doc.save(bio)
        return (
            bio.getvalue(),
            f"{kind}_report.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # PDF
    if not _HAS_REPORTLAB:
        txt = json.dumps(ctx, ensure_ascii=False, indent=2).encode("utf-8")
        return txt, f"{kind}_report_fallback.txt", "text/plain; charset=utf-8"

    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    w, h = A4
    y = h - 48
    c.setTitle(title)

    _register_cn_font()
    font = "Helvetica"
    if "Chinese" in pdfmetrics.getRegisteredFontNames():
        font = "Chinese"

    def line(txt: str, size: int = 11) -> None:
        nonlocal y
        c.setFont(font, size)
        for row in _wrap_for_canvas(txt, 42):
            if y < 56:
                c.showPage()
                c.setFont(font, size)
                y = h - 48
            c.drawString(40, y, row)
            y -= size + 4

    line(title, 16)
    line(f"时间: {ctx['generated_at']}", 10)
    line(f"建筑: {building_id or 'ALL'}", 10)
    line("", 10)
    line("时段汇总:", 12)
    line(json.dumps(ctx["period_summary"], ensure_ascii=False)[:2800], 9)
    line("", 8)
    line("指标与建议:", 12)
    line(json.dumps(ctx["indicators"], ensure_ascii=False), 9)
    for it in ctx["suggestions"]:
        line(f"- [{it.get('priority')}] {it.get('title')}", 9)
    c.save()
    return bio.getvalue(), f"{kind}_report.pdf", "application/pdf"


def _wrap_for_canvas(s: str, width: int) -> list[str]:
    if not s:
        return [""]
    out: list[str] = []
    cur = ""
    for ch in s:
        if len(cur) >= width and ch not in (" ", "，", "。"):
            out.append(cur)
            cur = ch
        else:
            cur += ch
    if cur:
        out.append(cur)
    return out if out else [""]


def _register_cn_font() -> None:
    if "Chinese" in pdfmetrics.getRegisteredFontNames():
        return
    candidates: list[tuple[str, int | None]] = [
        (r"C:\Windows\Fonts\msyh.ttc", 0),
        (r"C:\Windows\Fonts\simhei.ttf", None),
        ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", 0),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", 0),
    ]
    for path, sub in candidates:
        try:
            if sub is None:
                pdfmetrics.registerFont(TTFont("Chinese", path))
            else:
                pdfmetrics.registerFont(TTFont("Chinese", path, subfontIndex=sub))
            return
        except Exception:
            continue
